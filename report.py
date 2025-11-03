import os
import sys
import csv
import math
from datetime import datetime


def read_dataset_info(csv_path: str):
    """
    Reads basic info from a CSV dataset used in the project.
    - Returns: dict with n_rows, n_cols, class_counts (if 'Class' column exists), feature_names
    """
    info = {
        'n_rows': 0,
        'n_cols': 0,
        'class_counts': None,
        'feature_names': [],
        'dataset_name': os.path.basename(csv_path),
    }

    if not os.path.exists(csv_path):
        return info

    class_counts = {}
    with open(csv_path, 'r', newline='', encoding='utf-8') as f:
        reader = csv.reader(f)
        try:
            header = next(reader)
        except StopIteration:
            return info

        info['feature_names'] = header
        info['n_cols'] = len(header)
        class_index = None
        # Common label column names
        for cand in ('Class', 'class', 'label', 'target', 'y'):
            if cand in header:
                class_index = header.index(cand)
                break

        for row in reader:
            if not row:
                continue
            info['n_rows'] += 1
            if class_index is not None and class_index < len(row):
                cls = row[class_index]
                class_counts[cls] = class_counts.get(cls, 0) + 1

    if class_counts:
        info['class_counts'] = class_counts
    return info


def safe_imports():
    try:
        import matplotlib.pyplot as plt  # noqa: F401
        from docx import Document  # noqa: F401
        from docx.shared import Inches  # noqa: F401
    except Exception as e:
        return False, e
    return True, None


def plot_roc_auc(metrics_by_model, out_path):
    import matplotlib.pyplot as plt

    models = []
    values = []
    for name, data in metrics_by_model.items():
        v = data.get('metrics', {}).get('roc_auc', None)
        if v is not None and not (isinstance(v, float) and math.isnan(v)):
            models.append(name)
            values.append(v)

    if not models:
        return None

    plt.figure(figsize=(5, 3))
    bars = plt.bar(models, values, color=['#1f77b4', '#ff7f0e'])
    plt.ylim(0.0, 1.0)
    plt.title('Comparação de ROC AUC por Modelo')
    plt.ylabel('ROC AUC')
    for bar, val in zip(bars, values):
        plt.text(bar.get_x() + bar.get_width()/2.0, val + 0.01, f"{val:.4f}", ha='center', va='bottom', fontsize=9)
    plt.tight_layout()
    plt.savefig(out_path, dpi=160)
    plt.close()
    return out_path


def plot_metrics_grouped(metrics_by_model, out_path):
    import matplotlib.pyplot as plt

    metrics = ['accuracy', 'roc_auc', 'pr_auc']
    labels = {'accuracy': 'Acurácia', 'roc_auc': 'ROC AUC', 'pr_auc': 'PR AUC'}
    model_names = list(metrics_by_model.keys())

    data = {m: [] for m in metrics}
    for m in metrics:
        for name in model_names:
            val = metrics_by_model[name].get('metrics', {}).get(m, None)
            if val is None:
                data[m].append(float('nan'))
            else:
                data[m].append(val)

    x = range(len(model_names))
    width = 0.25

    plt.figure(figsize=(7, 4))
    for i, m in enumerate(metrics):
        xs = [xx + (i - 1) * width for xx in x]
        vals = data[m]
        # mask NaNs: plot only available values
        vals_to_plot = [v if (isinstance(v, float) and not math.isnan(v)) else 0.0 for v in vals]
        bar = plt.bar(xs, vals_to_plot, width=width, label=labels[m])
        for j, v in enumerate(vals):
            if isinstance(v, float) and not math.isnan(v):
                plt.text(xs[j], v + 0.01, f"{v:.4f}", ha='center', va='bottom', fontsize=8)

    plt.xticks(list(x), model_names)
    plt.ylim(0.0, 1.05)
    plt.ylabel('Valor da Métrica')
    plt.title('Métricas por Modelo')
    plt.legend()
    plt.tight_layout()
    plt.savefig(out_path, dpi=160)
    plt.close()
    return out_path


def plot_class_distribution(class_counts, out_path):
    import matplotlib.pyplot as plt

    if not class_counts:
        return None
    labels = list(class_counts.keys())
    values = [class_counts[k] for k in labels]
    plt.figure(figsize=(5, 3))
    bars = plt.bar(labels, values, color=['#2ca02c', '#d62728'])
    total = sum(values)
    for bar, val in zip(bars, values):
        pct = 100.0 * val / total if total else 0
        plt.text(bar.get_x() + bar.get_width()/2.0, val, f"{val} ({pct:.2f}%)", ha='center', va='bottom', fontsize=8)
    plt.title('Distribuição de Classes')
    plt.ylabel('Contagem')
    plt.tight_layout()
    plt.savefig(out_path, dpi=160)
    plt.close()
    return out_path


def generate_docx(output_path, dataset_info, metrics_by_model, notes=None):
    from docx import Document
    from docx.shared import Inches

    doc = Document()

    # Título
    doc.add_heading('Relatório de Modelos: MLP vs Gradient Boosting', 0)
    doc.add_paragraph(f"Data: {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    # Resumo executivo
    doc.add_heading('Resumo Executivo', level=1)
    doc.add_paragraph(
        'Este relatório compara dois modelos para detecção de fraudes: '
        'um Perceptron Multicamadas (MLP) customizado e um Gradient Boosting (GB) do scikit-learn. '
        'A comparação é feita com base nas métricas reportadas (Acurácia, ROC AUC e PR AUC, quando disponíveis).'
    )

    if notes:
        doc.add_paragraph(notes)

    # Dados do conjunto
    doc.add_heading('Visão Geral do Conjunto de Dados', level=1)
    ds_name = dataset_info.get('dataset_name')
    n_rows = dataset_info.get('n_rows')
    n_cols = dataset_info.get('n_cols')
    doc.add_paragraph(f"Arquivo: {ds_name}")
    if n_rows and n_cols:
        doc.add_paragraph(f"Instâncias: {n_rows} | Atributos (incluindo rótulo): {n_cols}")
    else:
        doc.add_paragraph('Não foi possível ler informações detalhadas do CSV informado.')

    if dataset_info.get('class_counts'):
        # Inserir gráfico de distribuição de classes
        class_chart = plot_class_distribution(dataset_info['class_counts'], 'chart_class_dist.png')
        if class_chart:
            doc.add_picture(class_chart, width=Inches(5.5))

    # Hiperparâmetros
    doc.add_heading('Hiperparâmetros', level=1)
    for name, data in metrics_by_model.items():
        doc.add_heading(name, level=2)
        params = data.get('params', {})
        if params:
            for k, v in params.items():
                doc.add_paragraph(f"{k}: {v}")
        else:
            doc.add_paragraph('Parâmetros não informados.')

    # Métricas
    doc.add_heading('Métricas de Avaliação', level=1)
    for name, data in metrics_by_model.items():
        doc.add_heading(name, level=2)
        mets = data.get('metrics', {})
        if not mets:
            doc.add_paragraph('Sem métricas informadas.')
        else:
            for k in ('accuracy', 'roc_auc', 'pr_auc'):
                if k in mets and mets[k] is not None and not (isinstance(mets[k], float) and math.isnan(mets[k])):
                    if k == 'accuracy':
                        doc.add_paragraph(f"Acurácia: {mets[k]:.4f}")
                    elif k == 'roc_auc':
                        doc.add_paragraph(f"ROC AUC: {mets[k]:.4f}")
                    elif k == 'pr_auc':
                        doc.add_paragraph(f"PR AUC: {mets[k]:.4f}")

    # Gráficos de comparação
    doc.add_heading('Gráficos de Comparação', level=1)
    roc_chart = plot_roc_auc(metrics_by_model, 'chart_roc_auc.png')
    if roc_chart:
        doc.add_paragraph('Comparação direta de ROC AUC entre os modelos:')
        doc.add_picture(roc_chart, width=Inches(5.5))

    metrics_chart = plot_metrics_grouped(metrics_by_model, 'chart_metrics.png')
    if metrics_chart:
        doc.add_paragraph('Métricas disponíveis por modelo (campos indisponíveis aparecem ausentes):')
        doc.add_picture(metrics_chart, width=Inches(5.5))

    # Análise
    doc.add_heading('Análise e Conclusões', level=1)
    analysis_p = doc.add_paragraph()
    analysis_p.add_run('Observações principais:')

    def fmt(v):
        try:
            return f"{float(v):.4f}" if not (isinstance(v, float) and math.isnan(v)) else 'N/D'
        except Exception:
            return 'N/D'

    mlp_mets = metrics_by_model.get('MLP', {}).get('metrics', {})
    gb_params = metrics_by_model.get('GB', {}).get('params', {})
    gb_mets = metrics_by_model.get('GB', {}).get('metrics', {})
    # Resultados resumidos
    gb_params_txt = (
        f"n_estimators={gb_params.get('n_estimators')}, "
        f"learning_rate={gb_params.get('learning_rate')}, "
        f"max_depth={gb_params.get('max_depth')}"
    )

    doc.add_paragraph(
        f"- Resultados obtidos (médias k-fold):\n"
        f"  • MLP: Acurácia={fmt(mlp_mets.get('accuracy'))}, ROC AUC={fmt(mlp_mets.get('roc_auc'))}, PR AUC={fmt(mlp_mets.get('pr_auc'))}.\n"
        f"  • Gradient Boosting ({gb_params_txt}): Acurácia={fmt(gb_mets.get('accuracy'))}, ROC AUC={fmt(gb_mets.get('roc_auc'))}, PR AUC={fmt(gb_mets.get('pr_auc'))}."
    )

    # Interpretação curta
    doc.add_paragraph(
        '- Interpretação: o Gradient Boosting apresentou ROC AUC e PR AUC superiores ao MLP, o que indica melhor capacidade de ranquear amostras por probabilidade de fraude e maior precisão/recall na região relevante. '
        'O MLP mostrou leve vantagem em acurácia (métrica sensível ao threshold), mas em problemas fortemente desbalanceados ROC AUC/PR AUC são métricas mais informativas.'
    )

    # Observação sobre pré-processamento e validação
    doc.add_paragraph(
        '- Observação sobre validação: os resultados aqui reportados foram obtidos usando validação cruzada estratificada e normalização Min-Max aplicada somente com estatísticas do conjunto de treino em cada fold (evitando vazamento de dados). '
        'Isso torna a comparação mais robusta e reduz o risco de sobrestimação do desempenho.'

    )

    # Comentários sobre PR AUC e desbalanceamento
    mlp_pr = mlp_mets.get('pr_auc')
    gb_pr = gb_mets.get('pr_auc')
    def is_missing(v):
        return v is None or (isinstance(v, float) and math.isnan(v))

    if is_missing(mlp_pr) or is_missing(gb_pr):
        doc.add_paragraph(
            '- Em cenários de detecção de fraude (desbalanceamento severo), PR AUC é especialmente informativa; '
            'recomenda-se garantir o cálculo de PR AUC para todos os modelos para uma comparação justa entre precisão/recall.'
        )
    else:
        doc.add_paragraph(
            f'- PR AUC observada (médias k-fold): Gradient Boosting = {fmt(gb_pr)}, MLP = {fmt(mlp_pr)}. '
            'No experimento, o GB apresentou PR AUC superior ao MLP, o que indica melhor precisão média nas regiões de recall de interesse. '
            'Em aplicações de fraude, isto pode significar menos falsos positivos para um nível de recall equivalente — útil quando o custo de investigações manuais é alto. '
            'Por outro lado, decisões operacionais devem ponderar custo de falsos negativos e preferências por recall vs. precisão.'
        )

    doc.add_paragraph(
        '- Considerando interpretabilidade e robustez, o GB costuma ser competitivo e pode oferecer melhor trade-off em ROC AUC; '
        'o MLP, por sua vez, tem bom potencial, mas depende sensivelmente de hiperparâmetros e normalização.'
    )

    doc.save(output_path)
    return output_path


def main():
    import argparse

    parser = argparse.ArgumentParser(description='Gera relatório DOCX comparando MLP vs GB')
    parser.add_argument('--dataset', type=str, default='creditcard_5k_80_20.csv', help='Caminho do CSV (para estatísticas e distribuição de classes)')
    parser.add_argument('--output', type=str, default='relatorio_modelos.docx', help='Arquivo DOCX de saída')
    parser.add_argument('--mlp-acc', type=float, default=0.9823)
    parser.add_argument('--mlp-roc', type=float, default=0.9746)
    parser.add_argument('--mlp-pr', type=float, default=0.9289)
    parser.add_argument('--gb-acc', type=float, default=0.9816)
    parser.add_argument('--gb-roc', type=float, default=0.9810)
    parser.add_argument('--gb-pr', type=float, default=0.9361)
    parser.add_argument('--gb-n-estimators', type=int, default=100)
    parser.add_argument('--gb-lr', type=float, default=0.05)
    parser.add_argument('--gb-max-depth', type=int, default=4)
    parser.add_argument('--note', type=str, default=(
        "Nota: Relatórios atualizados com os resultados de grid-search locais: "
        "MLP (n_oculta=16, taxa=0.2, epocas=20) e GB (n_estimators=100, lr=0.05, max_depth=4)."
    ))
    args = parser.parse_args()

    ok, err = safe_imports()
    if not ok:
        print('Dependências ausentes para gerar DOCX com gráficos.')
        print('Por favor, instale com:')
        print('  pip install python-docx matplotlib')
        print(f'Detalhes: {err}')
        sys.exit(2)

    dataset_info = read_dataset_info(args.dataset)

    metrics_by_model = {
        'MLP': {
            'params': {'n_oculta': 16, 'taxa_aprendizado': 0.2, 'epocas': 20},
            'metrics': {'accuracy': args.mlp_acc, 'roc_auc': args.mlp_roc, 'pr_auc': args.mlp_pr},
        },
        'GB': {
            'params': {'n_estimators': args.gb_n_estimators, 'learning_rate': args.gb_lr, 'max_depth': args.gb_max_depth},
            'metrics': {'accuracy': args.gb_acc, 'roc_auc': args.gb_roc, 'pr_auc': args.gb_pr},
        },
    }

    out = generate_docx(args.output, dataset_info, metrics_by_model, notes=args.note)
    print(f'Relatório gerado em: {out}')


if __name__ == '__main__':
    main()
